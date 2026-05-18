"""
Generate a branded Event Risk Report (PDF + DOCX) summarising the latest
scraped external event data and its impact on the Allianz exposure.

Reads from:
    allianz_silver.noaa_alerts
    allianz_silver.catastrophe_events
    allianz_silver.earthquake_events
    allianz_silver.weather_observations
    allianz_silver.news_events
    allianz_gold.event_risk_correlation
    allianz_gold.cat_event_pml

Writes to the UC bronze landing volume:
    /Volumes/<catalog>/allianz_bronze/landing/reports/event_risk_report_<ts>.pdf
    /Volumes/<catalog>/allianz_bronze/landing/reports/event_risk_report_<ts>.docx

Usage:
    uv run --no-project --with reportlab --with python-docx \
        --with "databricks-connect>=16.4,<17.0" \
        src/event_report.py
"""
from __future__ import annotations

import argparse
import os
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from databricks.connect import DatabricksSession

# --------------------------------------------------------------------------
# Config
# --------------------------------------------------------------------------
DEFAULT_CATALOG = "serverless_stable_xhky6g_catalog"
DEFAULT_VOLUME_SCHEMA = "allianz_bronze"
DEFAULT_VOLUME = "landing"
DEFAULT_PROFILE = "fe-vm-fevm-serverless-stable-xhky6g"

ALLIANZ_BLUE   = "#003781"
ALLIANZ_LIGHT  = "#006192"
ALLIANZ_ACCENT = "#FFAB00"
TEXT_GREY      = "#333333"


# --------------------------------------------------------------------------
# Data loaders
# --------------------------------------------------------------------------
def _safe(spark, sql: str, limit: int | None = None):
    """Run SQL and return list of dict rows; returns [] on any error."""
    try:
        df = spark.sql(sql if limit is None else f"{sql} LIMIT {limit}")
        return [r.asDict(recursive=True) for r in df.collect()]
    except Exception as e:
        print(f"  ! query failed: {e}")
        return []


def load_data(spark, catalog: str):
    out = {}

    out["counts"] = _safe(spark, f"""
        SELECT 'NOAA active alerts'  AS source, COUNT(*) AS n
            FROM {catalog}.allianz_silver.noaa_alerts
        UNION ALL SELECT 'GDACS catastrophes', COUNT(*)
            FROM {catalog}.allianz_silver.catastrophe_events
        UNION ALL SELECT 'USGS earthquakes',   COUNT(*)
            FROM {catalog}.allianz_silver.earthquake_events
        UNION ALL SELECT 'Weather obs',        COUNT(*)
            FROM {catalog}.allianz_silver.weather_observations
        UNION ALL SELECT 'News items',         COUNT(*)
            FROM {catalog}.allianz_silver.news_events
    """)

    out["alerts_by_severity"] = _safe(spark, f"""
        SELECT COALESCE(severity, 'Unknown') AS severity,
               COUNT(*) AS alerts,
               COUNT(DISTINCT event_type) AS distinct_event_types
        FROM {catalog}.allianz_silver.noaa_alerts
        GROUP BY severity ORDER BY alerts DESC
    """)

    out["top_exposed_alerts"] = _safe(spark, f"""
        SELECT event_type, severity, state_code, state_name,
               ROUND(exposed_tiv_usd / 1e9, 2) AS exposed_tiv_busd,
               exposed_policy_count, effective_utc, expires_utc
        FROM {catalog}.allianz_gold.event_risk_correlation
        WHERE exposed_tiv_usd IS NOT NULL
        ORDER BY exposed_tiv_usd DESC
    """, limit=20)

    out["top_pml_events"] = _safe(spark, f"""
        SELECT title, event_type, severity, state_code, event_time_utc,
               ROUND(state_tiv_usd / 1e9, 2)    AS state_tiv_busd,
               ROUND(severity_factor * 100, 1)  AS severity_factor_pct,
               ROUND(pml_estimate_usd / 1e6, 1) AS pml_musd
        FROM {catalog}.allianz_gold.cat_event_pml
        WHERE pml_estimate_usd IS NOT NULL
        ORDER BY pml_estimate_usd DESC
    """, limit=15)

    out["earthquakes"] = _safe(spark, f"""
        SELECT place, magnitude, depth_km, event_time_utc, severity
        FROM {catalog}.allianz_silver.earthquake_events
        ORDER BY magnitude DESC
    """, limit=10)

    out["weather_extremes"] = _safe(spark, f"""
        SELECT city, ROUND(temperature_f, 1) AS temp_f,
               ROUND(wind_speed_mph, 1) AS wind_mph,
               ROUND(wind_gust_mph, 1)  AS gust_mph,
               ROUND(precipitation_in, 2) AS precip_in,
               ROUND(humidity_pct, 0) AS humidity_pct,
               observed_time_utc
        FROM {catalog}.allianz_silver.weather_observations
        ORDER BY COALESCE(wind_gust_mph, 0) DESC
    """, limit=15)

    out["news"] = _safe(spark, f"""
        SELECT title, category, published_utc, link
        FROM {catalog}.allianz_silver.news_events
        ORDER BY published_utc DESC
    """, limit=12)

    return out


# --------------------------------------------------------------------------
# PDF generation (reportlab)
# --------------------------------------------------------------------------
def build_pdf(path: Path, run_ts: str, data: dict, catalog: str):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        PageBreak,
    )

    blue   = colors.HexColor(ALLIANZ_BLUE)
    light  = colors.HexColor(ALLIANZ_LIGHT)
    accent = colors.HexColor(ALLIANZ_ACCENT)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AllianzTitle", parent=styles["Title"], fontSize=26,
        textColor=blue, spaceAfter=12, alignment=0,
    )
    h2 = ParagraphStyle(
        "AllianzH2", parent=styles["Heading2"], fontSize=14,
        textColor=blue, spaceBefore=18, spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10,
        textColor=colors.HexColor(TEXT_GREY), spaceAfter=4,
    )
    meta = ParagraphStyle(
        "Meta", parent=styles["Normal"], fontSize=9,
        textColor=colors.HexColor("#666666"),
    )

    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch,  bottomMargin=0.6 * inch,
        title="Allianz Event Risk Report",
    )
    story = []

    # ── Cover ──────────────────────────────────────────────────────────
    story.append(Paragraph("ALLIANZ", ParagraphStyle(
        "Brand", parent=styles["Title"], fontSize=28, textColor=accent,
        spaceAfter=0)))
    story.append(Paragraph("Event Risk Report", title_style))
    story.append(Paragraph(
        f"Generated {run_ts} UTC · catalog {catalog}", meta))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Daily briefing on live weather alerts, catastrophe events, and "
        "humanitarian news scraped from public sources (NOAA, USGS, GDACS, "
        "Open-Meteo, ReliefWeb) — correlated with Allianz internal policy "
        "exposure to estimate potential impact.", body))

    # ── Executive Summary ──────────────────────────────────────────────
    story.append(Paragraph("Executive Summary", h2))
    rows = [["Source", "Records"]]
    for r in data.get("counts", []):
        rows.append([r["source"], f"{r['n']:,}"])
    story.append(_table(rows, [3.0 * inch, 2.0 * inch], blue, light))

    # ── NOAA by severity ───────────────────────────────────────────────
    if data.get("alerts_by_severity"):
        story.append(Paragraph("NOAA Active Alerts — by Severity", h2))
        rows = [["Severity", "Alerts", "Distinct event types"]]
        for r in data["alerts_by_severity"]:
            rows.append([str(r["severity"]),
                         f"{r['alerts']:,}",
                         f"{r['distinct_event_types']:,}"])
        story.append(_table(rows, [1.8*inch, 1.4*inch, 1.8*inch], blue, light))

    # ── Top exposed alerts ─────────────────────────────────────────────
    if data.get("top_exposed_alerts"):
        story.append(Paragraph(
            "Top NOAA Alerts × Allianz Exposure", h2))
        rows = [["Event", "Severity", "State", "Exposed TIV ($B)",
                 "Policies", "Effective (UTC)"]]
        for r in data["top_exposed_alerts"]:
            rows.append([
                _trim(r.get("event_type"), 28),
                str(r.get("severity") or ""),
                str(r.get("state_code") or ""),
                f"{r.get('exposed_tiv_busd', 0):,.2f}",
                f"{r.get('exposed_policy_count', 0):,}",
                _trim(str(r.get("effective_utc") or ""), 19),
            ])
        story.append(_table(
            rows, [1.7*inch, 0.9*inch, 0.6*inch, 1.2*inch, 0.8*inch, 1.4*inch],
            blue, light, font_size=8))

    # ── Catastrophe PML ───────────────────────────────────────────────
    story.append(PageBreak())
    if data.get("top_pml_events"):
        story.append(Paragraph("Catastrophe Events — PML Estimates", h2))
        rows = [["Event", "Type", "Sev", "State", "TIV ($B)",
                 "Sev factor", "PML ($M)"]]
        for r in data["top_pml_events"]:
            rows.append([
                _trim(r.get("title"), 30),
                str(r.get("event_type") or ""),
                str(r.get("severity") or ""),
                str(r.get("state_code") or ""),
                f"{r.get('state_tiv_busd', 0):,.2f}",
                f"{r.get('severity_factor_pct', 0):.1f}%",
                f"{r.get('pml_musd', 0):,.1f}",
            ])
        story.append(_table(
            rows, [1.8*inch, 1.0*inch, 0.5*inch, 0.6*inch, 1.0*inch,
                   0.9*inch, 0.9*inch],
            blue, light, font_size=8))

    # ── Earthquakes ────────────────────────────────────────────────────
    if data.get("earthquakes"):
        story.append(Paragraph("USGS Significant Earthquakes (last 30 days)", h2))
        rows = [["Location", "Magnitude", "Depth (km)", "Time (UTC)", "Severity"]]
        for r in data["earthquakes"]:
            rows.append([
                _trim(r.get("place"), 42),
                f"{r.get('magnitude') or '':}",
                f"{r.get('depth_km') or 0:.1f}",
                _trim(str(r.get("event_time_utc") or ""), 19),
                str(r.get("severity") or ""),
            ])
        story.append(_table(
            rows, [2.4*inch, 0.9*inch, 0.9*inch, 1.4*inch, 0.9*inch],
            blue, light, font_size=8))

    # ── Weather extremes ───────────────────────────────────────────────
    if data.get("weather_extremes"):
        story.append(PageBreak())
        story.append(Paragraph(
            "Weather Observations — CRESTA cities (sorted by gust)", h2))
        rows = [["City", "Temp (°F)", "Wind (mph)", "Gust (mph)",
                 "Precip (in)", "Hum %", "Time (UTC)"]]
        for r in data["weather_extremes"]:
            rows.append([
                _trim(r.get("city"), 20),
                f"{r.get('temp_f') or 0:.1f}",
                f"{r.get('wind_mph') or 0:.1f}",
                f"{r.get('gust_mph') or 0:.1f}",
                f"{r.get('precip_in') or 0:.2f}",
                f"{r.get('humidity_pct') or 0:.0f}",
                _trim(str(r.get("observed_time_utc") or ""), 16),
            ])
        story.append(_table(
            rows, [1.4*inch, 0.7*inch, 0.8*inch, 0.8*inch, 0.8*inch,
                   0.6*inch, 1.3*inch],
            blue, light, font_size=8))

    # ── News digest ────────────────────────────────────────────────────
    if data.get("news"):
        story.append(Paragraph("News Digest (ReliefWeb)", h2))
        for r in data["news"]:
            cat = r.get("category") or "General"
            story.append(Paragraph(
                f"<font color='{ALLIANZ_BLUE}'><b>[{cat}]</b></font> "
                f"{r.get('title', '')}", body))
            if r.get("link"):
                story.append(Paragraph(
                    f"<font color='#888888' size='8'>{r['link']}</font>", body))
            story.append(Spacer(1, 4))

    # ── Footer ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "<i>Generated by the Allianz Insurance Intelligence demo pipeline. "
        "Sources: NOAA NWS, USGS, GDACS, Open-Meteo, ReliefWeb. Internal "
        "exposure joined from Allianz silver/gold tables.</i>", meta))

    doc.build(story, onFirstPage=_branded_header, onLaterPages=_branded_header)


def _table(rows, col_widths, blue, light, font_size=9):
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), blue),
        ("TEXTCOLOR",    (0, 0), (-1, 0), colors.white),
        ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, -1), font_size),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                          [colors.white, colors.HexColor("#F2F6FB")]),
        ("GRID",         (0, 0), (-1, -1), 0.25, colors.HexColor("#DDDDDD")),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",        (1, 1), (-1, -1), "RIGHT"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _branded_header(canvas, doc):
    from reportlab.lib import colors
    canvas.saveState()
    canvas.setFillColor(colors.HexColor(ALLIANZ_BLUE))
    canvas.rect(0, doc.pagesize[1] - 0.25 * 72, doc.pagesize[0],
                0.25 * 72, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(0.6 * 72, doc.pagesize[1] - 0.18 * 72,
                      "ALLIANZ — Event Risk Report")
    canvas.drawRightString(doc.pagesize[0] - 0.6 * 72,
                           doc.pagesize[1] - 0.18 * 72,
                           f"page {doc.page}")
    canvas.restoreState()


def _trim(s, n):
    s = str(s) if s is not None else ""
    return s if len(s) <= n else s[: n - 1] + "…"


# --------------------------------------------------------------------------
# DOCX generation (python-docx)
# --------------------------------------------------------------------------
def build_docx(path: Path, run_ts: str, data: dict, catalog: str):
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ALLIANZ")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0xFF, 0xAB, 0x00)
    r.bold = True

    h = doc.add_heading("Event Risk Report", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x37, 0x81)

    doc.add_paragraph(f"Generated {run_ts} UTC · catalog {catalog}").italic = True
    doc.add_paragraph(
        "Daily briefing on live weather alerts, catastrophe events, and "
        "humanitarian news scraped from public sources (NOAA, USGS, GDACS, "
        "Open-Meteo, ReliefWeb) — correlated with Allianz internal policy "
        "exposure to estimate potential impact.")

    _add_table_section(doc, "Executive Summary",
        ["Source", "Records"],
        [(r["source"], f"{r['n']:,}") for r in data.get("counts", [])])

    _add_table_section(doc, "NOAA Active Alerts — by Severity",
        ["Severity", "Alerts", "Event types"],
        [(str(r["severity"]), f"{r['alerts']:,}", f"{r['distinct_event_types']:,}")
         for r in data.get("alerts_by_severity", [])])

    _add_table_section(doc, "Top NOAA Alerts × Allianz Exposure",
        ["Event", "Severity", "State", "Exposed TIV ($B)",
         "Policies", "Effective (UTC)"],
        [(_trim(r.get("event_type"), 28),
          str(r.get("severity") or ""),
          str(r.get("state_code") or ""),
          f"{r.get('exposed_tiv_busd', 0):,.2f}",
          f"{r.get('exposed_policy_count', 0):,}",
          _trim(str(r.get("effective_utc") or ""), 19))
         for r in data.get("top_exposed_alerts", [])])

    _add_table_section(doc, "Catastrophe Events — PML Estimates",
        ["Event", "Type", "Sev", "State", "TIV ($B)", "Sev factor", "PML ($M)"],
        [(_trim(r.get("title"), 30),
          str(r.get("event_type") or ""),
          str(r.get("severity") or ""),
          str(r.get("state_code") or ""),
          f"{r.get('state_tiv_busd', 0):,.2f}",
          f"{r.get('severity_factor_pct', 0):.1f}%",
          f"{r.get('pml_musd', 0):,.1f}")
         for r in data.get("top_pml_events", [])])

    _add_table_section(doc, "USGS Significant Earthquakes (last 30 days)",
        ["Location", "Magnitude", "Depth (km)", "Time (UTC)", "Severity"],
        [(_trim(r.get("place"), 42),
          str(r.get("magnitude") or ""),
          f"{r.get('depth_km') or 0:.1f}",
          _trim(str(r.get("event_time_utc") or ""), 19),
          str(r.get("severity") or ""))
         for r in data.get("earthquakes", [])])

    _add_table_section(doc, "Weather Observations (CRESTA cities)",
        ["City", "Temp (°F)", "Wind (mph)", "Gust (mph)",
         "Precip (in)", "Hum %", "Time (UTC)"],
        [(_trim(r.get("city"), 20),
          f"{r.get('temp_f') or 0:.1f}",
          f"{r.get('wind_mph') or 0:.1f}",
          f"{r.get('gust_mph') or 0:.1f}",
          f"{r.get('precip_in') or 0:.2f}",
          f"{r.get('humidity_pct') or 0:.0f}",
          _trim(str(r.get("observed_time_utc") or ""), 16))
         for r in data.get("weather_extremes", [])])

    doc.add_heading("News Digest (ReliefWeb)", level=2)
    for r in data.get("news", []):
        cat = r.get("category") or "General"
        p = doc.add_paragraph()
        run = p.add_run(f"[{cat}] ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x37, 0x81)
        p.add_run(r.get("title", ""))
        if r.get("link"):
            link_p = doc.add_paragraph()
            link_r = link_p.add_run(r["link"])
            link_r.font.size = Pt(8)
            link_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    foot = doc.add_paragraph()
    foot_r = foot.add_run(
        "Generated by the Allianz Insurance Intelligence demo pipeline. "
        "Sources: NOAA NWS, USGS, GDACS, Open-Meteo, ReliefWeb. Internal "
        "exposure joined from Allianz silver/gold tables.")
    foot_r.italic = True
    foot_r.font.size = Pt(8)
    foot_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(str(path))


def _add_table_section(doc, heading, header_row, body_rows):
    from docx.shared import RGBColor, Pt
    h = doc.add_heading(heading, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x37, 0x81)

    if not body_rows:
        doc.add_paragraph("(no data)").italic = True
        return

    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_row))
    table.style = "Light Grid Accent 1"
    for j, col in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(body_rows, start=1):
        for j, v in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(v)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


# --------------------------------------------------------------------------
# Upload — uses the SDK Files API so it works inside a Databricks job
# (the `databricks` CLI is not installed in serverless task runtimes).
# --------------------------------------------------------------------------
def upload(local_path: Path, remote: str, profile: str):
    try:
        from databricks.sdk import WorkspaceClient
        # Inside a Databricks job, WorkspaceClient() picks up the runtime
        # credentials automatically. Locally, fall back to the named profile.
        try:
            w = WorkspaceClient()
        except Exception:
            w = WorkspaceClient(profile=profile)
        with open(local_path, "rb") as fh:
            w.files.upload(file_path=remote, contents=fh, overwrite=True)
        return True
    except Exception as e:
        print(f"  ✗ upload failed: {e}")
        return False


# --------------------------------------------------------------------------
# Driver
# --------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--catalog", default=os.environ.get("ALLIANZ_CATALOG", DEFAULT_CATALOG))
    parser.add_argument("--volume-schema", default=DEFAULT_VOLUME_SCHEMA)
    parser.add_argument("--volume", default=DEFAULT_VOLUME)
    parser.add_argument("--profile", default=os.environ.get(
        "DATABRICKS_CONFIG_PROFILE", DEFAULT_PROFILE))
    parser.add_argument("--format", choices=["pdf", "docx", "both"], default="both")
    parser.add_argument("--local-only", action="store_true",
                        help="Write to ./output/reports instead of UC volume.")
    args = parser.parse_args()

    run_ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    print("=" * 70)
    print("ALLIANZ EVENT RISK REPORT")
    print("=" * 70)
    print(f"Catalog : {args.catalog}")
    print(f"Format  : {args.format}")
    print(f"Run TS  : {run_ts}")
    print()

    os.environ["DATABRICKS_CONFIG_PROFILE"] = args.profile
    spark = (DatabricksSession.builder
             .profile(args.profile).serverless().getOrCreate())
    print("Loading event data from silver/gold…")
    data = load_data(spark, args.catalog)
    spark.stop()

    formats = ["pdf", "docx"] if args.format == "both" else [args.format]
    local_dir = Path("output/reports") if args.local_only else Path(tempfile.gettempdir())
    local_dir.mkdir(parents=True, exist_ok=True)

    for fmt in formats:
        local = local_dir / f"event_risk_report_{run_ts}.{fmt}"
        print(f"\nBuilding {fmt.upper()}…")
        if fmt == "pdf":
            build_pdf(local, run_ts, data, args.catalog)
        else:
            build_docx(local, run_ts, data, args.catalog)
        size_kb = local.stat().st_size / 1024
        print(f"  built {local}  ({size_kb:.1f} KB)")

        if args.local_only:
            print(f"  ✓ local only")
            continue

        remote = (f"/Volumes/{args.catalog}/{args.volume_schema}/"
                  f"{args.volume}/reports/event_risk_report_{run_ts}.{fmt}")
        if upload(local, remote, args.profile):
            print(f"  → {remote}")
        if not args.local_only:
            local.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
